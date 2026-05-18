from pathlib import Path
from tempfile import NamedTemporaryFile

import pandas as pd
from bs4 import BeautifulSoup
from docx import Document
from fastapi import UploadFile
from pypdf import PdfReader

from sdlc_copilot.models import SourceDocument


SUPPORTED_SUFFIXES = {".pdf", ".txt", ".md", ".html", ".htm", ".docx", ".csv", ".json"}


async def load_upload(file: UploadFile) -> SourceDocument:
    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        raise ValueError(f"Unsupported file type '{suffix}'. Supported: {sorted(SUPPORTED_SUFFIXES)}")

    raw = await file.read()
    if not raw:
        raise ValueError("Uploaded file is empty.")

    with NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(raw)
        tmp_path = Path(tmp.name)

    try:
        text = load_path(tmp_path)
    finally:
        tmp_path.unlink(missing_ok=True)

    return SourceDocument(
        filename=file.filename or "upload",
        content_type=file.content_type,
        text=text,
        metadata={"source": "upload", "size_bytes": len(raw)},
    )


def load_path(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _load_pdf(path)
    if suffix in {".txt", ".md"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix in {".html", ".htm"}:
        return BeautifulSoup(path.read_text(encoding="utf-8", errors="ignore"), "html.parser").get_text("\n")
    if suffix == ".docx":
        return "\n".join(p.text for p in Document(path).paragraphs)
    if suffix == ".csv":
        return pd.read_csv(path).to_markdown(index=False)
    if suffix == ".json":
        return pd.read_json(path).to_json(orient="records", indent=2)
    raise ValueError(f"Unsupported file type '{suffix}'.")


def _load_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        pages.append(f"\n\n[Page {index}]\n{page_text}")
    text = "\n".join(pages).strip()
    if not text:
        raise ValueError("PDF contained no extractable text. OCR is required for scanned documents.")
    return text
